#!/usr/bin/env python3
"""
audit_template_colors.py — read Joe's colour convention back out of the
templates and check the field maps against it.

JOE'S CONVENTION (stated 2026-07-31): in a CARR template, BLUE or RED text is
text that gets replaced, coloured so a human eye can see what still needs
answering. A letter that goes out is all black.

WHY THAT IS A VERIFICATION GIFT. The colour is the template author's own list of
slots, written years before anyone tried to automate the thing. ORDER 13 derived
its three field maps by reading the templates and proposing a map. Comparing the
proposed slots against the coloured runs checks that derivation against an
independent source rather than against itself:

  COLOURED BUT UNMAPPED — the template says a human replaces this and no map
      slot covers it. A slot the analysis missed.
  MAPPED BUT UNCOLOURED — a map slot on text the template never marked. Either
      a real slot the author never coloured, or over-mapping.

NOTHING IS ADDED TO A MAP HERE. The maps are Joe-reviewed territory and the
whole point of `unreviewed:true` is that a human reads them before a document
fills from them. This writes findings into each map's `review_note` so they are
in front of whoever does that review, and it changes nothing else.

  .venv/bin/python fill-engine/audit_template_colors.py            # report
  .venv/bin/python fill-engine/audit_template_colors.py --write    # + review_note
"""
from __future__ import annotations

import argparse
import json
import os
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
sys.path.insert(0, os.path.dirname(HERE))
from fill_document import colored_runs  # noqa: E402
from lib.drive_recovery import add_recovery_arguments, require_recovery  # noqa: E402

MAPS = os.path.join(HERE, "field-maps")

TEMPLATES = {
    "loi-lease": "DNA/Deal Management/Templates/RFP-LOI-Lease Template.docx",
    "loi-purchase": "DNA/Deal Management/Templates/LOI-Purchase.docx",
}


def sendable_note(slug: str) -> str:
    """Which file leaves, per template kind. Joe's doctrine of 2026-07-31,
    superseding the blanket 'clients get PDFs, never working files' the factory
    was built on."""
    if slug == "loi-grid":
        return ("SENDABLE FORMAT (Joe's doctrine, 2026-07-31, superseding the old blanket 'clients "
                "get PDFs, never working files'): this is a SPREADSHEET template, so the PDF is the "
                "sendable artifact and the working workbook never leaves. Concealing the formulas "
                "is the whole point. Naming a sendable format is not permission to send: Joe sends, "
                "and no verb in this system can.")
    return ("SENDABLE FORMAT (Joe's doctrine, 2026-07-31, superseding the old blanket 'clients get "
            "PDFs, never working files'): this is a LETTER template, so the WORKING .docx is what "
            "goes to the listing agent. They edit and revise it, and that revision IS the "
            "negotiation; the PDF is the record and preview copy. Because the .docx is outbound it "
            "is SCRUBBED before it routes, of comments, tracked changes, comment authors, custom "
            "properties and the authorship metadata inherited from CARR's corporate letterhead. "
            "Naming a sendable format is not permission to send: Joe sends, and no verb can.")


def apply_sendable_note(slug: str) -> str:
    path = os.path.join(MAPS, f"{slug}.json")
    with open(path) as fh:
        m = json.load(fh)
    note = sendable_note(slug)
    if note not in (m.get("review_note") or ""):
        m["review_note"] = ((m.get("review_note", "") + "\n\n") if m.get("review_note") else "") + note
        with open(path, "w") as fh:
            json.dump(m, fh, indent=2)
            fh.write("\n")
    return path


def control_addresses(path: str) -> dict[str, str]:
    """Slots the template marks with a Word CONTENT CONTROL instead of a colour.

    THE SECOND MARKER, and finding it is what turned a noisy result into a
    usable one. The first run of this audit reported the lease template's letter
    date and contact line as "mapped but uncoloured", which reads as
    over-mapping. They are neither: they are black text inside content controls,
    a different way of saying the same thing ("a human fills this"). So the
    convention has two markers and only one of them is a colour. An address that
    carries a control is reported in its own class rather than lumped in with
    genuinely unmarked text, because the two demand opposite answers from a
    reviewer.
    """
    import docx
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    d = docx.Document(path)
    hits: dict[str, str] = {}
    for pi, p in enumerate(d.paragraphs):
        if p._p.findall(f".//{W}sdt"):
            hits[f"para:{pi}"] = p.text.strip()[:70]
    for ti, t in enumerate(d.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                if cell._tc.findall(f".//{W}sdt"):
                    hits[f"table:{ti}:{ri}:{ci}"] = cell.text.strip()[:70]
    return hits


def audit(slug: str, rel: str, vault: str) -> dict:
    path = os.path.join(vault, rel)
    runs = colored_runs(path)
    controls = control_addresses(path)
    with open(os.path.join(MAPS, f"{slug}.json")) as fh:
        m = json.load(fh)
    slots = m.get("slots", {})
    mapped = {s["where"] for s in slots.values() if "where" in s}
    # A repeat block addresses many cells through one declaration; its own
    # addressing lives inside the block, so it is counted as covering them.
    for s in slots.values():
        for a in s.get("covers", []):
            mapped.add(a)

    by_addr: dict[str, list[str]] = {}
    for r in runs:
        by_addr.setdefault(r["where"], []).append(f"{r['color']} {r['text'].strip()[:70]!r}")

    colored_unmapped = {a: v for a, v in by_addr.items() if a not in mapped}
    unmarked = sorted(a for a in mapped if a not in by_addr and a not in controls)
    by_control = {a: controls[a] for a in sorted(mapped) if a not in by_addr and a in controls}
    return {"slug": slug, "template": rel, "colored_addresses": len(by_addr),
            "colored_runs": len(runs), "mapped_slots": len(mapped),
            "colored_unmapped": colored_unmapped, "mapped_uncolored": unmarked,
            "mapped_by_control": by_control, "controls": len(controls),
            "map_path": os.path.join(MAPS, f"{slug}.json")}


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--write", action="store_true",
                    help="append the finding to each map's review_note (adds no slot)")
    add_recovery_arguments(ap)
    a = ap.parse_args()
    try:
        vault = str(require_recovery(
            a, "document-template API for color audit"))
    except ValueError as exc:
        print(f"audit-template-colors: STOP: {exc}", file=sys.stderr)
        return 2

    for slug, rel in TEMPLATES.items():
        r = audit(slug, rel, vault)
        print(f"\n=== {slug} — {os.path.basename(rel)}")
        print(f"    {r['colored_runs']} coloured runs across {r['colored_addresses']} addresses · "
              f"{r['mapped_slots']} mapped slot addresses")
        print(f"    COLOURED BUT UNMAPPED ({len(r['colored_unmapped'])}) — the template says a human "
              f"replaces these and no slot covers them:")
        for addr, texts in sorted(r["colored_unmapped"].items()):
            print(f"      {addr:<18} {texts[0]}" + (f"   (+{len(texts) - 1} more runs)"
                                                    if len(texts) > 1 else ""))
        print(f"    MAPPED, MARKED BY A CONTENT CONTROL ({len(r['mapped_by_control'])}) — the "
              f"template's OTHER way of saying 'a human fills this'. Correctly mapped:")
        for addr, txt in r["mapped_by_control"].items():
            print(f"      {addr:<18} {txt!r}")
        print(f"    MAPPED BUT UNMARKED ({len(r['mapped_uncolored'])}) — black text the template "
              f"marks neither way. Real slot the author never marked, or over-mapping:")
        for addr in r["mapped_uncolored"]:
            print(f"      {addr}")

        if a.write:
            with open(r["map_path"]) as fh:
                m = json.load(fh)
            note = (
                f"COLOUR CROSS-CHECK (2026-07-31, ORDER 20 fold-in). Joe's convention: blue or red "
                f"marks text that gets replaced; a letter that goes out is all black. Checked this "
                f"map against the template's own coloured runs, which are the author's independent "
                f"list of slots. {r['colored_runs']} coloured runs across {r['colored_addresses']} "
                f"addresses. COLOURED BUT UNMAPPED ({len(r['colored_unmapped'])}): "
                + ("; ".join(f"{k} {v[0]}" for k, v in sorted(r['colored_unmapped'].items())) or "none")
                + f". MAPPED AND MARKED BY A CONTENT CONTROL instead of a colour, correctly mapped "
                  f"({len(r['mapped_by_control'])}): "
                + (", ".join(r["mapped_by_control"]) or "none")
                + f". MAPPED BUT UNMARKED EITHER WAY ({len(r['mapped_uncolored'])}): "
                + (", ".join(r["mapped_uncolored"]) or "none")
                + ". NOTHING WAS ADDED TO THIS MAP: the map is reviewed territory, and a slot found "
                  "by colour is a candidate for the human reading it, not a change made on its behalf."
                + "\n\n" + sendable_note(r["slug"]))
            m["review_note"] = (m.get("review_note", "") + ("\n\n" if m.get("review_note") else "")
                                + note) if note not in (m.get("review_note") or "") else m["review_note"]
            with open(r["map_path"], "w") as fh:
                json.dump(m, fh, indent=2)
                fh.write("\n")
            print(f"    review_note updated: {r['map_path']}")

    if a.write:
        # The grid map has no coloured runs to survey (it is an xlsx), but it
        # carries the OTHER half of the sendable ruling and a reviewer reading
        # only one map should still find it.
        print(f"\n=== loi-grid — sendable note only (xlsx: no coloured-run survey applies)")
        print(f"    review_note updated: {apply_sendable_note('loi-grid')}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

#!/usr/bin/env python3
"""
fill_document.py — the generalized fill engine (ORDER 13(b), wave2-design 2c).

The vault's `DNA/Deal Management/fill-engine/fill_deal_template.py` filled ONE
shape: an xlsx by {cell: value} plus a LibreOffice recalc. This generalizes it
to the two shapes the document factory needs — docx slot filling and xlsx cell
filling — with a PDF render for either, and it never modifies a template in
place: every fill copies first.

  fill(template_path, out_path, edits)        -> out_path      (docx or xlsx)
  to_pdf(src_path, out_pdf_path)              -> out_pdf_path

`edits` is a list of {"where": <address>, "text": <string>}:

  docx   para:<i>                 the i-th body paragraph
         table:<t>:<r>:<c>        table t, row r, cell c
  xlsx   cell:<Sheet>!<A1>        a named sheet
         cell:<A1>                the active sheet

WHY TEXT-ONLY EDITS: branding is everything this engine must not break. A CARR
template carries its letterhead in header1.xml, its palette in theme1.xml and
its table styling in the table properties — none of which is touched here. The
engine rewrites the TEXT of an addressed run and deletes its siblings, so the
filled cell keeps the template's own font, size, weight and colour. Everything
not named in `edits` is byte-identical to the template, which is the property
ORDER 13 asks for and the reason no template is ever edited in place.

LibreOffice is required for the PDF render and for the xlsx recalc (the same
dependency the vault engine already had). It is NOT bundled: `soffice_path()`
looks in the three places it lives on a Mac and raises a plain-language error
naming the install command if it finds none.
"""
from __future__ import annotations

import os
import re
import shutil
import subprocess
import tempfile

SOFFICE_CANDIDATES = [
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/opt/homebrew/bin/soffice",
    "/usr/local/bin/soffice",
    "/usr/bin/soffice",
]


class FillError(Exception):
    pass


def soffice_path() -> str:
    env = os.environ.get("CARR_SOFFICE")
    if env and os.path.exists(env):
        return env
    for p in SOFFICE_CANDIDATES:
        if os.path.exists(p):
            return p
    found = shutil.which("soffice") or shutil.which("libreoffice")
    if found:
        return found
    raise FillError(
        "LibreOffice not found. The fill engine needs it for the PDF render and "
        "the xlsx recalc. Install it with:  brew install --cask libreoffice  "
        "(or set CARR_SOFFICE to the soffice binary)."
    )


# ---------------------------------------------------------------- addressing

_PARA = re.compile(r"^para:(\d+)$")
_CELL = re.compile(r"^table:(\d+):(\d+):(\d+)$")
_XL = re.compile(r"^cell:(?:(.+)!)?([A-Z]{1,3}\d{1,7})$")


def parse_address(where: str):
    m = _PARA.match(where)
    if m:
        return ("para", int(m.group(1)))
    m = _CELL.match(where)
    if m:
        return ("table", int(m.group(1)), int(m.group(2)), int(m.group(3)))
    m = _XL.match(where)
    if m:
        return ("cell", m.group(1), m.group(2))
    raise FillError(f"unrecognised slot address {where!r} "
                    "(expected para:N, table:T:R:C, or cell:[Sheet!]A1)")


# ---------------------------------------------------------------- docx fill

def _unwrap_content_controls(el) -> None:
    """Replace every Word content control (w:sdt) in `el` with its own content.

    CARR's letterhead templates are built on content controls bound to a custom
    XML part: the lease LOI has one for the letter Date and two more for Broker
    Email and Broker Phone. python-docx does NOT see runs inside a control, so
    a naive fill leaves them standing: the first C-112 draft rendered with the
    template's stale 'October 27, 2023' and with 'first.last@carr.us' plus a
    live 'Click or tap here to enter text' placeholder sitting beside the value
    that was supposed to replace them. Found in the rendered PDF, not in the
    XML — which is the fill-engine workflow's own rule about reading the
    rendered document rather than a text extraction, earning its keep.

    Unwrapping is the correct semantic: a filled document is a document, not a
    form. The control's runs are kept exactly as they are, so the character
    formatting the control carried survives.
    """
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    while True:
        sdts = el.findall(f".//{W}sdt")
        if not sdts:
            return
        # deepest last: process one at a time and re-scan, so nesting is safe
        sdt = sdts[-1]
        parent = sdt.getparent()
        content = sdt.find(f"{W}sdtContent")
        idx = list(parent).index(sdt)
        if content is not None:
            for child in list(content):
                parent.insert(idx, child)
                idx += 1
        parent.remove(sdt)


# ------------------------------------------------- Joe's colour convention
# JOE'S OWN RULE FOR THE CARR TEMPLATES (stated 2026-07-31): text in BLUE or RED
# is text that gets replaced, coloured "for human eyes to know the text gets
# replaced". A letter that actually goes out is ALL BLACK.
#
# The engine inherits the addressed run's formatting on purpose, which is what
# keeps the letterhead intact — and it is also what made every filled value come
# out in the placeholder's blue or red. Measured on the C-112 draft before this
# change: the tenant name, the signature block, the $17.00 rate and the 3%
# escalation all rendered RED, in a document whose whole point is that it is
# ready for Joe to read. So a fill now STATES the colour rather than inheriting
# it, and the two states say opposite things at a glance:
#   BLACK — this value is real and final.
#   RED   — this is an OWED marker. A human still has to answer it.
# Red for owed is not a new convention. It is Joe's existing one, pointed at the
# one thing left in a draft that still needs him.
BLACK = "000000"
RED = "FF0000"


def _set_run_color(run, rgb: str | None) -> None:
    """Force a run's colour, clearing any theme colour that would override it.

    Setting `w:color/@w:val` alone is not enough: a run carrying `w:themeColor`
    keeps taking its colour from the theme, so the attribute is removed rather
    than left to fight with the value.
    """
    if rgb is None:
        return
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    rPr = run._element.get_or_add_rPr()
    c = rPr.find(f"{W}color")
    if c is None:
        c = rPr.makeelement(f"{W}color", {})
        rPr.append(c)
    c.set(f"{W}val", rgb)
    for attr in (f"{W}themeColor", f"{W}themeTint", f"{W}themeShade"):
        if c.get(attr) is not None:
            del c.attrib[attr]


def _set_paragraph_text(para, text: str, color: str | None = None) -> None:
    """Replace a paragraph's text, keeping the FIRST run's formatting.

    Word stores a paragraph as a sequence of runs; a template's placeholder is
    often split across several of them by spell-check or tracked-change history.
    Writing into run 0 and dropping the rest is the only way to guarantee the
    result reads as one clean value in the template's own typeface. A paragraph
    with no runs (an empty line) gets one created, inheriting the paragraph
    style rather than the document default.

    `color` overrides the inherited font colour and nothing else: size, weight,
    typeface and the paragraph's own styling all still come from the template.
    """
    _unwrap_content_controls(para._p)
    runs = para.runs
    if not runs:
        r = para.add_run(text)
        _set_run_color(r, color)
        return
    runs[0].text = text
    _set_run_color(runs[0], color)
    for extra in runs[1:]:
        extra._element.getparent().remove(extra._element)


def _set_cell_text(cell, text: str, color: str | None = None) -> None:
    """Replace a table cell's text across the whole cell, first paragraph wins.

    Multi-line template cells (the LOI's Broker Commission cell is three
    sentences over one paragraph) collapse to the single filled value; extra
    paragraphs are removed so a stale second line cannot survive a fill.
    """
    _unwrap_content_controls(cell._tc)
    paras = cell.paragraphs
    _set_paragraph_text(paras[0], text, color)
    for extra in paras[1:]:
        extra._element.getparent().remove(extra._element)


def edit_color(e: dict) -> str | None:
    """The colour one edit writes in. Owed markers red, real values black.

    An explicit `color` on the edit wins, so a caller that needs something else
    is not fighting the default; absent that, the edit's own `owed` flag decides.
    """
    if "color" in e:
        v = e["color"]
        return None if v in (None, "inherit") else {"black": BLACK, "red": RED}.get(v, v)
    return RED if e.get("owed") else BLACK


def fill_docx(template: str, out: str, edits: list[dict]) -> str:
    import docx  # imported lazily so an xlsx-only run needs no python-docx

    if os.path.abspath(template) == os.path.abspath(out):
        raise FillError("refusing to fill a template in place; give a distinct out path")
    d = docx.Document(template)
    for e in edits:
        addr = parse_address(e["where"])
        text = "" if e.get("text") is None else str(e["text"])
        color = edit_color(e)
        if addr[0] == "para":
            i = addr[1]
            if i >= len(d.paragraphs):
                raise FillError(f"{e['where']}: document has {len(d.paragraphs)} paragraphs")
            _set_paragraph_text(d.paragraphs[i], text, color)
        elif addr[0] == "table":
            _, t, r, c = addr
            if t >= len(d.tables):
                raise FillError(f"{e['where']}: document has {len(d.tables)} tables")
            tb = d.tables[t]
            if r >= len(tb.rows) or c >= len(tb.columns):
                raise FillError(f"{e['where']}: table {t} is {len(tb.rows)}x{len(tb.columns)}")
            _set_cell_text(tb.rows[r].cells[c], text, color)
        else:
            raise FillError(f"{e['where']}: cell addresses belong to xlsx templates")
    os.makedirs(os.path.dirname(os.path.abspath(out)), exist_ok=True)
    d.save(out)
    return out


# ------------------------------------------------ the outbound docx scrub
# JOE'S DOCTRINE, stated 2026-07-31 and it corrects what the factory inherited:
# an LOI letter goes to the LISTING AGENT AS A .docx, because the counterparty
# edits and revises it and that IS the negotiation workflow. The PDF is the
# record and preview copy. Spreadsheets are the other way round: the grid goes
# out as a PDF precisely so nobody sees the formulas, and the working xlsx never
# leaves. So a .docx from this factory is an OUTBOUND artifact, and an outbound
# Word file must read as a clean CARR document.
#
# IT DOES NOT, OUT OF THE BOX, AND THAT IS A MEASURED FINDING RATHER THAN A
# PRECAUTION. The C-112 draft filed in Joe's deal folder on 2026-07-31 carried,
# inherited straight from CARR's corporate letterhead template:
#     dc:creator          Andrew Mondy
#     cp:lastModifiedBy   Mike Jorgenson
#     cp:revision         23
#     created / modified  2018-01-24 / 2023-10-27
#     app.xml Template    C:\Users\andrew\AppData\Roaming\Microsoft\Templates\...
# A named individual's Windows user directory, on a document about to be handed
# to the other side of a negotiation. Nothing about it is secret and nothing
# about it is intentional either, which is the definition of a fingerprint.
#
# WHAT IS SCRUBBED AND WHAT IS DELIBERATELY NOT. Removed: comments and their
# parts, tracked changes (insertions accepted, deletions dropped, revision-only
# formatting records removed), comment authors, custom document properties, and
# the authorship and revision metadata above. KEPT: the letterhead, the theme,
# the styles, the numbering, every byte of branding. The scrub is metadata
# surgery, never content surgery, and the sha256 check on the TEMPLATE still
# holds because the template is never what gets scrubbed.
CARR_AUTHOR = "CARR"

_DROP_PARTS = ("word/comments.xml", "word/commentsExtended.xml", "word/commentsIds.xml",
               "word/commentsExtensible.xml", "word/people.xml", "docProps/custom.xml")


def _strip_revisions(body, W: str) -> dict:
    """Accept insertions, drop deletions, remove revision-only records."""
    found = {"insertions_accepted": 0, "deletions_removed": 0, "format_revisions_removed": 0,
             "comment_marks_removed": 0, "moves_resolved": 0}

    for tag in (f"{W}ins", f"{W}moveTo"):
        for el in body.findall(f".//{tag}"):
            parent = el.getparent()
            idx = list(parent).index(el)
            for child in list(el):
                parent.insert(idx, child)
                idx += 1
            parent.remove(el)
            found["insertions_accepted" if tag.endswith("ins") else "moves_resolved"] += 1

    for tag in (f"{W}del", f"{W}moveFrom"):
        for el in body.findall(f".//{tag}"):
            el.getparent().remove(el)
            found["deletions_removed"] += 1

    for tag in (f"{W}rPrChange", f"{W}pPrChange", f"{W}sectPrChange", f"{W}tblPrChange",
                f"{W}trPrChange", f"{W}tcPrChange", f"{W}cellIns", f"{W}cellDel"):
        for el in body.findall(f".//{tag}"):
            el.getparent().remove(el)
            found["format_revisions_removed"] += 1

    for tag in (f"{W}commentRangeStart", f"{W}commentRangeEnd", f"{W}commentReference"):
        for el in body.findall(f".//{tag}"):
            el.getparent().remove(el)
            found["comment_marks_removed"] += 1

    return found


def scrub_docx(path: str, author: str = CARR_AUTHOR) -> dict:
    """Make one .docx safe to hand to a counterparty. Returns what it found.

    Reports rather than whispers: every value it clears is named in the return,
    because "what was in this file before we cleaned it" is the finding, and a
    scrub nobody can read is indistinguishable from a scrub that did nothing.
    """
    import zipfile
    import docx
    from lxml import etree

    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    found: dict = {"path": path}

    d = docx.Document(path)
    found.update(_strip_revisions(d.element.body, W))

    cp = d.core_properties
    found["metadata_cleared"] = {
        "author": cp.author, "last_modified_by": cp.last_modified_by,
        "revision": cp.revision, "comments": cp.comments, "category": cp.category,
        "keywords": cp.keywords, "subject": cp.subject,
        "created": str(cp.created) if cp.created else None,
        "modified": str(cp.modified) if cp.modified else None}
    cp.author = author
    cp.last_modified_by = author
    cp.revision = 1
    cp.comments = ""
    cp.category = ""
    cp.keywords = ""
    cp.subject = ""
    found["metadata_cleared"]["last_printed"] = (str(cp.last_printed) if cp.last_printed else None)
    # python-docx refuses None for lastPrinted, so the element goes rather than
    # the value: "this document was printed on <date>" is exactly the kind of
    # history an outbound file has no reason to carry.
    for el in list(cp._element):
        if etree.QName(el).localname == "lastPrinted":
            cp._element.remove(el)
    d.save(path)

    # Part-level removal. python-docx leaves an unreferenced comments part in
    # the package, and an orphan part is still a part: the text is recoverable
    # by anyone who opens the zip. So the package is rewritten with those
    # entries gone and the two manifests that name them repaired.
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        entries = {n: z.read(n) for n in names}
    dropped = [n for n in names if n in _DROP_PARTS]
    found["parts_removed"] = dropped
    found["app_properties_cleared"] = {}

    if "docProps/app.xml" in entries:
        root = etree.fromstring(entries["docProps/app.xml"])
        for local in ("Template", "Manager", "Company", "TotalTime", "LastAuthor", "HyperlinkBase"):
            for el in root.iter():
                if etree.QName(el).localname == local and (el.text or "").strip():
                    found["app_properties_cleared"][local] = el.text
                    el.text = "CARR" if local == "Company" else ""
        entries["docProps/app.xml"] = etree.tostring(root, xml_declaration=True,
                                                     encoding="UTF-8", standalone=True)

    if dropped:
        ct = etree.fromstring(entries["[Content_Types].xml"])
        for el in list(ct):
            part = (el.get("PartName") or "").lstrip("/")
            if part in dropped:
                ct.remove(el)
        entries["[Content_Types].xml"] = etree.tostring(ct, xml_declaration=True,
                                                        encoding="UTF-8", standalone=True)
        for rels in ("word/_rels/document.xml.rels", "_rels/.rels"):
            if rels not in entries:
                continue
            base = os.path.dirname(os.path.dirname(rels)) or ""
            rr = etree.fromstring(entries[rels])
            for el in list(rr):
                tgt = (el.get("Target") or "").lstrip("/")
                full = os.path.normpath(os.path.join(base, tgt)) if base else tgt
                if full.replace(os.sep, "/") in dropped or tgt in dropped:
                    rr.remove(el)
            entries[rels] = etree.tostring(rr, xml_declaration=True,
                                           encoding="UTF-8", standalone=True)

    tmp = path + ".scrub"
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as z:
        for n in names:
            if n in dropped:
                continue
            z.writestr(n, entries[n])
    os.replace(tmp, path)

    # Left alone on purpose, and named so nobody thinks it was missed:
    # customXml/ parts are the content-control bindings the template ships with.
    # The fill unwraps every control, so they bind nothing; they carry no
    # authorship and removing them safely means rewriting three more manifests
    # for no gain in what a counterparty can learn.
    found["kept_deliberately"] = [n for n in names if n.startswith(("customXml/", "word/glossary/"))]
    return found


# --------------------------------------------------------- the finish rules
# ORDER 22, and it exists because ORDER 20's fold-in 1 was only half a rule.
# Filled values were forced black and OWED markers red, which left the CARRIED
# text — CARR's own standing language, sitting in the template in the template's
# replace-me blue and red — coloured on a document that goes to a listing agent.
#
# JOE'S CORRECTION (2026-07-31, now an active shared rule) settles which side
# gives: every LOI term is negotiable and a template value is a placeholder
# holding the most common case, so the colour in a template means "review this
# per deal", not "this is CARR policy". The template authorship is therefore
# CORRECT and must not be forked. The document is what changes: on the sendable
# render every run is black except the markers that still owe Joe an answer.
#
# The verification surface moves with the rule. A reader can no longer tell a
# carried run from a filled one by its ink, so `audit_template_colors.py` and the
# `carried` list in the run's own report are what say what was carried, and the
# colour left in the file means one thing only: a human still owes an answer.
OWED_PREFIX = re.compile(r"^\s*\[OWED:")


def _force_run_color(r_el, rgb: str) -> None:
    """Force one w:r element's colour, clearing any theme colour beside it.

    Works on the raw element rather than on a python-docx Run so it reaches runs
    inside a content control, which `paragraph.runs` cannot see. The colour
    element is placed through python-docx's own rPr accessor so it lands in the
    schema's required child order instead of merely at the end.
    """
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    c = r_el.get_or_add_rPr().get_or_add_color()
    c.set(f"{W}val", rgb)
    for attr in (f"{W}themeColor", f"{W}themeTint", f"{W}themeShade"):
        if c.get(attr) is not None:
            del c.attrib[attr]


def _run_text(r_el) -> str:
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    return "".join(t.text or "" for t in r_el.findall(f".//{W}t"))


def finish_colors(path: str) -> dict:
    """Force every body run BLACK except the engine's own OWED markers, red.

    Body only: header1.xml holds the letterhead and is never touched, so the
    branding keeps its own colours and the .emf and the theme stay byte-identical
    to the template. Empty runs are skipped because a run with no text has no ink
    to correct, and touching it would change the package for no visible gain.
    """
    import docx

    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    d = docx.Document(path)
    out: dict = {"forced_black": 0, "owed_kept_red": 0, "recolored": []}
    for r in d.element.body.findall(f".//{W}r"):
        text = _run_text(r)
        if not text.strip():
            continue
        rPr = r.find(f"{W}rPr")
        c = None if rPr is None else rPr.find(f"{W}color")
        prev = "" if c is None else (c.get(f"{W}val") or "").upper()
        theme = None if c is None else c.get(f"{W}themeColor")
        was_colored = bool(theme) or prev not in ("", "AUTO", "000000")
        if OWED_PREFIX.match(text):
            _force_run_color(r, RED)
            out["owed_kept_red"] += 1
            continue
        _force_run_color(r, BLACK)
        out["forced_black"] += 1
        if was_colored:
            out["recolored"].append({"was": prev or f"theme:{theme}", "text": text[:90]})
    d.save(path)
    return out


def has_unresolved_option(text: str) -> bool:
    """A pipe is the template author's way of writing 'pick one'."""
    return " | " in text or text.rstrip().endswith("|")


def resolve_unresolved_options(path: str, labels: dict | None = None) -> dict:
    """Turn any surviving 'A | B' into a red OWED marker naming the alternatives.

    ORDER 22(b). A pipe that reaches the finish is not text, it is a question
    nobody answered, and a letter reaching the listing agent with both
    alternatives still in it is a worse failure than a colour (measured on the
    C-112 draft of 2026-07-31: the Broker Commission, HVAC and Electrical rows
    all shipped their pipes). The alternatives are named as the template writes
    them rather than paraphrased, so the answer Joe gives is a choice between the
    template's own words.
    """
    import docx

    labels = labels or {}
    d = docx.Document(path)
    found: list[dict] = []

    def handle(where: str, text: str, setter) -> None:
        if not has_unresolved_option(text):
            return
        segs = [s.strip() for s in re.split(r"\s*\|\s*", text) if s.strip()]
        label = labels.get(where) or "unresolved template option"
        marker = (f"[OWED: {label}. The template offers alternatives here and nobody has "
                  "chosen: " + " OR ".join(f'"{s}"' for s in segs) + "]")
        setter(marker)
        found.append({"where": where, "was": text, "alternatives": segs, "marker": marker})

    for pi, p in enumerate(d.paragraphs):
        handle(f"para:{pi}", p.text, lambda t, p=p: _set_paragraph_text(p, t, RED))
    for ti, t in enumerate(d.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                handle(f"table:{ti}:{ri}:{ci}", cell.text,
                       lambda t, cell=cell: _set_cell_text(cell, t, RED))
    if found:
        d.save(path)
    return {"count": len(found), "resolved_to_owed": found}


def document_pipes(path: str) -> list[str]:
    """Every body run still carrying a pipe. The done-test reads this, not a report."""
    import docx

    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    d = docx.Document(path)
    return [_run_text(r) for r in d.element.body.findall(f".//{W}r") if "|" in _run_text(r)]


# ------------------------------------------------------- the colour audit
# Joe's convention read back out of a file. Two uses, one function: the finish
# check on a produced draft (anything still coloured is either an unfilled slot
# or a template patch nobody made), and the template survey that tells a field
# map which slots the template itself says exist.

def colored_runs(path: str, ignore_black: bool = True) -> list[dict]:
    """Every run in a docx carrying a non-black explicit colour.

    Returns [{where, color, text}] with `where` in the engine's own address
    grammar, so a finding names a slot address a map can be compared against.
    """
    import docx
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    d = docx.Document(path)
    # THE SAME BLIND SPOT ORDER 13 FOUND, and it bites here too: python-docx
    # cannot see runs inside a Word content control, so a scan that trusts
    # `paragraph.runs` reports a coloured control as uncoloured. The first run of
    # this audit called the lease template's letter date and contact line
    # "mapped but uncoloured" for exactly that reason, which would have read as
    # over-mapping when it was a scanner that could not see. Unwrapping happens
    # on the IN-MEMORY document only; nothing is written back, and the file on
    # disk is never touched by an audit.
    _unwrap_content_controls(d.element.body)
    out = []

    def scan(where, para):
        for r in para.runs:
            rPr = r._element.find(f"{W}rPr")
            c = None if rPr is None else rPr.find(f"{W}color")
            if c is None:
                continue
            val = (c.get(f"{W}val") or "").upper()
            theme = c.get(f"{W}themeColor")
            if ignore_black and val in ("", "AUTO", "000000") and not theme:
                continue
            if not r.text.strip():
                continue
            out.append({"where": where, "color": val or f"theme:{theme}", "text": r.text})

    for pi, p in enumerate(d.paragraphs):
        scan(f"para:{pi}", p)
    for ti, t in enumerate(d.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    scan(f"table:{ti}:{ri}:{ci}", p)
    return out


# ---------------------------------------------------------------- xlsx fill

def fill_xlsx(template: str, out: str, edits: list[dict], recalc: bool = True) -> str:
    import openpyxl

    if os.path.abspath(template) == os.path.abspath(out):
        raise FillError("refusing to fill a template in place; give a distinct out path")
    wb = openpyxl.load_workbook(template)
    for e in edits:
        addr = parse_address(e["where"])
        if addr[0] != "cell":
            raise FillError(f"{e['where']}: paragraph/table addresses belong to docx templates")
        _, sheet, ref = addr
        ws = wb[sheet] if sheet else wb.active
        ws[ref] = e.get("text")
    wb.calculation.fullCalcOnLoad = True
    os.makedirs(os.path.dirname(os.path.abspath(out)), exist_ok=True)
    if not recalc:
        wb.save(out)
        return out
    # LibreOffice recalc, same trick the vault engine used: save, round-trip
    # through headless Calc so every formula holds a computed value the moment
    # the client opens it, then move the result into place.
    indir = tempfile.mkdtemp(prefix="carrfill_in_")
    outdir = tempfile.mkdtemp(prefix="carrfill_out_")
    try:
        src = os.path.join(indir, "book.xlsx")
        wb.save(src)
        subprocess.run(
            [soffice_path(), "--headless", "--calc", "--convert-to",
             "xlsx:Calc MS Excel 2007 XML", "--outdir", outdir, src],
            check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
            env={**os.environ, "HOME": tempfile.gettempdir()})
        produced = os.path.join(outdir, "book.xlsx")
        if not os.path.exists(produced):
            raise FillError("LibreOffice produced no recalculated workbook")
        shutil.move(produced, out)
    finally:
        shutil.rmtree(indir, ignore_errors=True)
        shutil.rmtree(outdir, ignore_errors=True)
    return out


def fill(template: str, out: str, edits: list[dict]) -> str:
    ext = os.path.splitext(template)[1].lower()
    if ext == ".docx":
        return fill_docx(template, out, edits)
    if ext in (".xlsx", ".xlsm"):
        return fill_xlsx(template, out, edits)
    raise FillError(f"no filler for {ext} templates")


# ---------------------------------------------------------------- pdf render

def to_pdf(src: str, out_pdf: str) -> str:
    """Headless LibreOffice render. The client-facing artifact is ALWAYS the PDF
    (tool-contracts §3: clients get PDFs, never working files) so this is not an
    optional step — a document with no PDF is not finished."""
    outdir = tempfile.mkdtemp(prefix="carrpdf_")
    try:
        subprocess.run(
            [soffice_path(), "--headless", "--convert-to", "pdf", "--outdir", outdir, src],
            check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
            env={**os.environ, "HOME": tempfile.gettempdir()})
        produced = os.path.join(outdir, os.path.splitext(os.path.basename(src))[0] + ".pdf")
        if not os.path.exists(produced):
            raise FillError(f"LibreOffice produced no PDF for {src}")
        os.makedirs(os.path.dirname(os.path.abspath(out_pdf)), exist_ok=True)
        shutil.move(produced, out_pdf)
    finally:
        shutil.rmtree(outdir, ignore_errors=True)
    return out_pdf


if __name__ == "__main__":
    import json
    import sys
    if len(sys.argv) < 4:
        print("usage: fill_document.py <template> <out> <edits.json> [--pdf <out.pdf>]")
        raise SystemExit(2)
    fill(sys.argv[1], sys.argv[2], json.load(open(sys.argv[3])))
    print("wrote", sys.argv[2])
    if "--pdf" in sys.argv:
        p = to_pdf(sys.argv[2], sys.argv[sys.argv.index("--pdf") + 1])
        print("wrote", p)

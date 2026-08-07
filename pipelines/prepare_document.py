#!/usr/bin/env python3
"""
prepare_document.py — the LOCAL half of the document factory (ORDER 13).

THE SPLIT, and why it is not a shortcut. `prepare-document` is an MCP verb in a
Cloudflare Worker: no filesystem, no LibreOffice, no OneDrive. So the verb owns
everything the RECORD layer can answer — resolve the deal, apply the reviewed
field map, decide every slot's value or mark it owed, create the `document` row
under the envelope with its event — and returns a fill PLAN. This script owns
everything only a machine standing in Joe's file system can do: copy the
template, write the slots, render the PDF, run the leak guard and the writing
lint, and file the result. Neither half re-derives the other's work; the plan is
the whole contract between them.

  # 1. get the plan (deployed Worker, or local-verb.mjs against a branch)
  # 2. produce the files:
  ./run.sh ... | .venv/bin/python pipelines/prepare_document.py plan.json
       [--route onedrive|staging] [--finish]

  --route staging   (default) writes into out/documents/ only
  --route onedrive  files working + PDF into the deal's OneDrive folder, unless
                    the writing lint returns a HARD finding — a HARD finding
                    blocks the OneDrive copy by design and says so
  --finish          with DATABASE_URL set, uploads the working file and the PDF
                    to the R2 archive under the self-enforced quota, writes their
                    attachment rows, and patches the document row (lint/leak
                    results, attachment pointers)
  --r2-dry-run      exercise the quota gate and report, upload nothing, write
                    nothing
  --no-client-copy  suppress the CLIENT COPY artifact (ORDER 23(c)), which is
                    otherwise produced whenever the field map names a row the
                    client must not see

WHICH FILE IS THE SENDABLE ONE (Joe's doctrine, 2026-07-31, superseding the old
blanket "clients get PDFs, never working files"): a LETTER goes to the listing
agent as the working .docx, because the counterparty edits and revises it and
that editing IS the negotiation; the PDF is the record and preview copy. A
SPREADSHEET goes out as the PDF so the formulas stay ours, and the working
workbook never leaves. Naming a sendable format is not permission to send:
Joe sends, and no verb in this system can.

THE FINISH RULES (ORDER 22, from Joe's correction that every LOI term is
negotiable and a template value is a placeholder holding the most common case,
never CARR policy). On the sendable render: every run is BLACK except the OWED
markers, which are RED, so the only colour a reader sees is a question somebody
still owes an answer to. What the template carried is read off the run's own
`carried` list and off audit_template_colors.py, not off the document's ink. And
any text still carrying a pipe is an UNRESOLVED CHOICE, never text: it becomes a
red OWED marker naming the template's own alternatives. Both are checked on the
produced file, and that check BLOCKS the OneDrive copy the way a HARD lint
finding does.

ROWS THAT LEAVE THE LETTER (ORDER 23, from Joe's field-map review). Two of his
rulings, both carried in the reviewed field map as data and evaluated here at
finish: a `drop_when` row is deleted when the deal makes the term meaningless (an
NNN lease has no base year), and an `audience` row is deleted from the CLIENT
COPY only (the commission ask goes to the listing agent, never to the client).
Drops run AFTER every fill because deleting a row renumbers the rows below it,
and a row is never dropped on a guess: an unresolved condition keeps the row.
The client copy carries CLIENT-COPY in its filename and passes the same gates as
the working file, so nothing client-named can be mistaken for the counterparty's
copy and no artifact leaves here ungated.

Every outbound .docx is SCRUBBED before it routes — comments, tracked changes,
comment authors, custom properties, and the authorship metadata inherited from
CARR's corporate template (the C-112 draft carried a named individual's Windows
template path until this ran). Branding is untouched; see fill_document.py.

THE QUOTA (ORDER 20, Joe's requirement 2026-07-31) is a HARD cap, not an alert.
It lives in `system_config` under `r2.quota_gb` and DEFAULTS TO 8 GB when that
row is absent, which is under Cloudflare's 10 GB free tier on purpose. An upload
that would cross it is REFUSED: nothing uploads, nothing is deleted to make
room, the OneDrive copies stand, and the archive copy is recorded as owed on the
document row with the reason. See lib/r2_archive.py.

NOTHING IS EVER SENT. The output is a DRAFT for Joe to read.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
from typing import Any

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(REPO, "fill-engine"))
sys.path.insert(0, os.path.join(REPO, "lib"))
from fill_document import (fill, to_pdf, FillError, colored_runs, scrub_docx,  # noqa: E402
                           finish_colors, resolve_unresolved_options, document_pipes,
                           drop_rows, parse_row_address, table_row_texts)
import r2_archive as r2  # noqa: E402

VAULT = os.environ.get(
    "CARR_VAULT",
    "/Users/booko/Library/CloudStorage/GoogleDrive-joe.bookout.carr.us@gmail.com/My Drive/CARR AI")
ONEDRIVE_DEALS = os.environ.get(
    "CARR_ONEDRIVE_DEALS",
    "/Users/booko/Library/CloudStorage/OneDrive-CARR,Inc/Joe's Folder/Deals/Active Deals")
STAGING = os.path.join(REPO, "out", "documents")

# The leak guard's named checks. Each is a fact about CARR's own material that
# must never reach a client-facing file, and each fires on a signature rather
# than on a judgment call.
LEAK_PATTERNS = [
    (r"205-?643-?6555", "the placeholder phone number (standing data rule: never stored, never shipped)"),
    (r"INTERNAL[\s_-]?ONLY", "an internal-only marker"),
    (r"CommissionComparison|BaseRentCalculator", "the commission calculator, which never reaches a client"),
    (r"sf_commission_placeholder|sf_close_date_placeholder", "a Salesforce placeholder column name"),
]


def doc_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        import docx
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for r in t.rows:
                parts.extend(c.text for c in r.cells)
        return "\n".join(x for x in parts if x and x.strip())
    if ext in (".xlsx", ".xlsm"):
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for c in row:
                    if isinstance(c.value, str) and c.value.strip():
                        parts.append(c.value)
        return "\n".join(parts)
    return ""


def leak_guard(text: str, listing_side_names: list[str]) -> list[str]:
    findings = []
    for pat, why in LEAK_PATTERNS:
        if re.search(pat, text, re.I):
            findings.append(f"{why} (pattern {pat})")
    for name in listing_side_names:
        if name and re.search(re.escape(name), text, re.I):
            findings.append(f"a listing-side party name appears in a client-facing file: {name}")
    return findings


OWED_MARKER = re.compile(r"^\s*\[OWED:")


def color_check(path: str) -> dict:
    """The no-coloured-text finish check, now a GATE (ORDER 22(a)).

    Joe's convention: blue or red in a CARR template means "this gets replaced",
    and a letter that goes out is all black. ORDER 20 forced filled values black
    and OWED markers red but left the CARRIED runs in the template's own
    replace-me colours, and reported them rather than blocking — because
    blocking on carried text would have stopped every LOI this factory ever
    produces.

    That trade is gone: the finish now forces every run black except the OWED
    markers, so nothing legitimate is left coloured and the check can hold a real
    line. The invariant is arithmetic rather than judgment — every coloured run
    in the sendable file is an OWED marker, and every OWED marker is red — so a
    failure means the finish rule did not reach some run, which is a defect in
    this engine and not a fact about the deal.

    A pipe is checked separately and over the whole body, because a run inside a
    content control is invisible to a paragraph scan and " | " reaching the
    listing agent is worse than any colour.
    """
    if os.path.splitext(path)[1].lower() != ".docx":
        return {"applies": False, "passed": True}
    runs = colored_runs(path)
    markers = [r for r in runs if OWED_MARKER.match(r["text"])]
    not_marker = [r for r in runs if not OWED_MARKER.match(r["text"])]
    not_red = [r for r in markers if r["color"] != "FF0000"]
    pipes = document_pipes(path)
    return {"applies": True,
            "colored_runs": len(runs), "owed_markers": len(markers),
            "colored_but_not_owed": not_marker, "owed_but_not_red": not_red,
            "pipes_remaining": pipes,
            "passed": not not_marker and not not_red and not pipes}


def slot_labels(plan: dict) -> dict:
    """Address -> the label a human knows the slot by, for an OWED marker's text."""
    out = {}
    for group in ("owed", "carried", "partial"):
        for s in plan.get(group, []) or []:
            if s.get("where"):
                out[s["where"]] = s.get("label") or s.get("slot") or ""
    for e in plan.get("edits", []) or []:
        out.setdefault(e["where"], e.get("slot", ""))
    return {k: v for k, v in out.items() if v}


# ------------------------------------------------------------- row drops
# ORDER 23, from Joe's field-map review. Two rules, one operation, and the field
# map carries both as DATA rather than this file carrying them as code:
#
#   drop_when  a term that does not exist on this deal. base_year drops when the
#              lease is triple-net, because an NNN lease has no base year — the
#              row is about nothing, which is not the same as unanswered.
#   audience   a term that exists and is not this reader's business. The broker
#              commission goes to the listing agent and never to the client, so
#              it drops from the CLIENT COPY and stays in the working .docx.
#
# The map is Joe's reviewed territory and the conditions are his words, so a
# condition changes by editing the map, never by editing this file. Both are
# evaluated at FINISH, after every fill: a drop renumbers the rows below it, so
# dropping first would move addresses out from under the map that wrote them.

FIELD_MAPS = os.path.join(REPO, "fill-engine", "field-maps")


def load_field_map(plan: dict) -> dict:
    """The reviewed map for this plan's template, or an empty map if none is here.

    An absent map is not an error: the plan is the contract between the Worker
    and this script, and the map is read only for the two finish rules the
    record layer cannot evaluate. No map means no drops, said out loud in the
    report rather than assumed.
    """
    slug = (plan.get("template") or {}).get("slug") or ""
    path = os.path.join(FIELD_MAPS, f"{slug}.json")
    if not slug or not os.path.exists(path):
        return {"_missing": path, "slots": {}}
    with open(path) as fh:
        m = json.load(fh)
    m["_path"] = path
    return m


def resolved_slot_values(plan: dict) -> dict:
    """slot -> the text actually written for it. Owed markers are NOT values."""
    out = {}
    for e in plan.get("edits", []) or []:
        if e.get("slot") and not e.get("owed"):
            out[e["slot"]] = str(e.get("text") or "")
    return out


def conditional_drops(fmap: dict, plan: dict) -> dict:
    """Evaluate every `drop_when` in the map against this deal's resolved slots.

    A row is never dropped on a guess: if the governing slot is owed or missing,
    the row stays and the report says why. Both outcomes are returned, because
    "the base year row is still here" is a fact Joe should be able to read
    without re-deriving the condition himself.
    """
    vals = resolved_slot_values(plan)
    drop, kept = [], []
    for name, s in (fmap.get("slots") or {}).items():
        cond = s.get("drop_when")
        if not cond:
            continue
        gov, want = cond.get("slot"), cond.get("resolves_to") or []
        have = vals.get(gov)
        row = {"slot": name, "label": s.get("label"), "where": s.get("where"),
               "governed_by": gov, "resolves_to": want, "resolved_value": have,
               "why": cond.get("why", "")}
        if have is None:
            if cond.get("when_unresolved") == "drop":
                row["decision"] = f"DROPPED: {gov} is unresolved and the map says drop anyway"
                drop.append(row)
            else:
                row["decision"] = (f"KEPT: {gov} is unresolved on this deal, and a row is never "
                                   "dropped on a guess")
                kept.append(row)
        elif have in want:
            row["decision"] = f"DROPPED: {gov} resolved to {have!r}"
            drop.append(row)
        else:
            row["decision"] = f"KEPT: {gov} resolved to {have!r}, which is not a drop value"
            kept.append(row)
    return {"drop": drop, "kept": kept}


def audience_drops(fmap: dict, audience: str = "client") -> list[dict]:
    """Every slot the map says this audience must not see."""
    out = []
    for name, s in (fmap.get("slots") or {}).items():
        aud = s.get("audience") or {}
        if audience in (aud.get("drop_for") or []):
            out.append({"slot": name, "label": s.get("label"), "where": s.get("where"),
                        "why": aud.get("why", "")})
    return out


def shift_for_prior_drops(where: str, prior: list[str]) -> str:
    """Re-address a row after earlier rows in the same table were removed.

    THE BUG THIS EXISTS TO PREVENT, stated because it is invisible in a test that
    only drops one row: the client copy is made from the FINISHED working file,
    which on an NNN lease has already lost row 9. The commission row the map
    addresses as table:0:18 is row 17 in that file. Address arithmetic is the
    honest fix — the map keeps naming the template's own numbering, and every
    consumer of a modified document adjusts for what it already removed. The
    caller verifies the result by reading the row's own label before deleting it,
    because arithmetic that is wrong should fail loudly on the label, not quietly
    on a neighbouring row.
    """
    t, r = parse_row_address(where)
    gone = sum(1 for p in prior if parse_row_address(p)[0] == t and parse_row_address(p)[1] < r)
    return f"table:{t}:{r - gone}"


def run_lint(text: str, label: str) -> dict:
    """tools/writing-lint.py as the gate it already is. Exit 1 = HARD finding."""
    tmp = os.path.join(STAGING, f".lint-{label}.txt")
    os.makedirs(STAGING, exist_ok=True)
    with open(tmp, "w") as fh:
        fh.write(text)
    p = subprocess.run(
        [sys.executable, os.path.join(REPO, "tools", "writing-lint.py"), tmp, "--surface", "proposal"],
        capture_output=True, text=True)
    os.remove(tmp)
    return {"hard": p.returncode != 0, "exit": p.returncode,
            "report": (p.stdout or "") + (p.stderr or "")}


def sha256_of(path: str) -> tuple[str, int]:
    h = hashlib.sha256()
    n = 0
    with open(path, "rb") as fh:
        for chunk in iter(lambda: fh.read(65536), b""):
            h.update(chunk)
            n += len(chunk)
    return h.hexdigest(), n


def find_deal_folder(plan: dict) -> str | None:
    """Match Joe's EXISTING per-deal folder rather than minting a new one.

    His convention is one folder per deal under Active Deals, named the way he
    thinks of the deal, which is not always the client party name (C-112's folder
    is 'Gulf Coast Pelvic Health'; the party row says 'Gulf Coast Pelvic Floor').
    Matching by shared words beats matching by string equality, and a miss
    returns None so the caller reports it instead of creating a second folder
    beside the real one.
    """
    if not os.path.isdir(ONEDRIVE_DEALS):
        return None
    wanted = {plan["deal"].get("name"), plan["deal"].get("client_name"),
              plan["deal"].get("org_name")}
    tokens = set()
    for w in wanted:
        if w:
            tokens |= {t.lower() for t in re.findall(r"[A-Za-z]+", w) if len(t) > 2}
    best, best_score = None, 0
    for entry in os.listdir(ONEDRIVE_DEALS):
        full = os.path.join(ONEDRIVE_DEALS, entry)
        if not os.path.isdir(full) or entry.startswith("_") or entry.startswith("."):
            continue
        et = {t.lower() for t in re.findall(r"[A-Za-z]+", entry) if len(t) > 2}
        score = len(tokens & et)
        if score > best_score:
            best, best_score = full, score
    return best if best_score >= 2 else None


def build_client_copy(working: str, plan: dict, fmap: dict, prior_drops: list[str]) -> dict:
    """The third artifact (ORDER 23(c)): the finished letter minus what the client must not see.

    Rendered FROM the finished sendable rather than from a second fill, so it is
    the same document Joe already read with rows removed — there is no second
    path for a value to differ down. It carries CLIENT-COPY in its own filename
    for one reason: nothing named for the client can be handed to a counterparty
    by accident, and a name is the only thing that survives being dragged into an
    email. The listing-agent working copy keeps every row; the record PDF is the
    full version. This changes no send semantics — Joe sends, nothing here can.

    The label check is the safety rail. A row is only deleted after this function
    has read that row's own label cell and found the slot's label in it, so an
    address that has drifted stops the client copy instead of silently deleting
    the wrong term.
    """
    drops = audience_drops(fmap, "client")
    base, ext = os.path.splitext(working)
    out: dict[str, Any] = {"audience": "client", "produced": False,
                           "declared_drops": drops,
                           "docx": base + "-CLIENT-COPY" + ext, "pdf": base + "-CLIENT-COPY.pdf"}
    if fmap.get("_missing"):
        out["why_not"] = f"no field map on disk at {fmap['_missing']}, so no audience rule to apply"
        return out
    if not drops:
        out["why_not"] = "the field map declares no audience drops for a client copy"
        return out

    rows = table_row_texts(working, 0)
    wheres, reasons, mismatches = [], {}, []
    for d in drops:
        if not d.get("where"):
            mismatches.append({"slot": d["slot"], "problem": "the map slot has no address"})
            continue
        adj = shift_for_prior_drops(d["where"], prior_drops)
        t, r = parse_row_address(adj)
        label = (d.get("label") or "").strip().lower()
        actual = rows[r][0].strip() if t == 0 and r < len(rows) and rows[r] else ""
        if label and label not in actual.lower():
            mismatches.append({"slot": d["slot"], "map_address": d["where"],
                               "address_in_finished_file": adj,
                               "expected_label": d.get("label"), "row_label_found": actual})
            continue
        wheres.append(adj)
        reasons[adj] = f"{d.get('label')} — {d.get('why', '')}"
    if mismatches:
        out["why_not"] = ("STOPPED: a row addressed for the client copy does not carry the label "
                          "the map names, so no row was deleted. See address_mismatches.")
        out["address_mismatches"] = mismatches
        return out

    shutil.copy2(working, out["docx"])
    out["dropped"] = drop_rows(out["docx"], wheres, reasons)
    out["scrub"] = scrub_docx(out["docx"])
    to_pdf(out["docx"], out["pdf"])                # LibreOffice render = the structural check
    text = doc_text(out["docx"])
    out["leak_findings"] = leak_guard(text, plan.get("listing_side_names", []))
    lint = run_lint(text, (plan["basename"][:32] + "-client"))
    out["lint_hard"], out["lint_report"] = lint["hard"], lint["report"]
    out["color_gate"] = color_check(out["docx"])
    out["docx_sha256"], out["docx_bytes"] = sha256_of(out["docx"])
    out["pdf_sha256"], out["pdf_bytes"] = sha256_of(out["pdf"])
    out["produced"] = True
    out["passed"] = (out["color_gate"]["passed"] and not out["lint_hard"]
                     and not out["leak_findings"])
    out["human_gate"] = ("The client copy is a DRAFT for Joe. It exists so the commission ask is "
                         "not in front of the client; it is not a second sendable and nothing here "
                         "sends anything.")
    return out


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("plan", help="the prepare-document plan JSON ('-' for stdin)")
    ap.add_argument("--route", choices=["staging", "onedrive"], default="staging")
    ap.add_argument("--finish", action="store_true",
                    help="archive to R2 + write attachment rows + patch the document row "
                         "(needs DATABASE_URL)")
    ap.add_argument("--r2-bucket", default="carr-documents")
    ap.add_argument("--r2-dry-run", action="store_true",
                    help="run the quota gate and report, upload nothing, write no rows")
    ap.add_argument("--no-client-copy", action="store_true",
                    help="skip the CLIENT-COPY artifact even when the field map declares "
                         "audience drops (ORDER 23(c)); the client copy is produced by default "
                         "for any docx whose map names a client-audience drop")
    a = ap.parse_args()

    plan = json.load(sys.stdin if a.plan == "-" else open(a.plan))
    tmpl_rel = plan["template"]["source_path"]
    tmpl = os.path.join(VAULT, tmpl_rel)
    if not os.path.exists(tmpl):
        print(f"STOP: template not found at {tmpl}", file=sys.stderr)
        return 2
    ext = os.path.splitext(tmpl)[1].lower()
    os.makedirs(STAGING, exist_ok=True)
    working = os.path.join(STAGING, plan["basename"] + ext)
    pdf = os.path.join(STAGING, plan["basename"] + ".pdf")

    # ---- fill. The template is opened read-only and copied; nothing writes to it.
    before_sha, _ = sha256_of(tmpl)
    # The `owed` flag rides through to the engine on purpose: it is what decides
    # the colour Joe reads the draft by (black = real, red = still needs him).
    edits = [{"where": e["where"], "text": e["text"], "owed": bool(e.get("owed"))}
             for e in plan["edits"]]
    # ---- the finish rules (ORDER 22), between the fill and the render so the
    # PDF is a render of the finished file rather than of an intermediate one.
    # Order matters: pipes become OWED markers FIRST, so the colour pass sees
    # them as markers and leaves them red instead of blackening a live question.
    fmap = load_field_map(plan)
    options, color_finish, row_drops = None, None, None
    prior_drops: list[str] = []
    try:
        fill(tmpl, working, edits)
        if ext == ".docx":
            options = resolve_unresolved_options(working, slot_labels(plan))
            color_finish = finish_colors(working)
            # ORDER 23(a)+(b), LAST of the three finish rules on purpose: a drop
            # renumbers every row below it, so it runs after every address the
            # map wrote has already been resolved against the template's own
            # numbering.
            cond = conditional_drops(fmap, plan)
            prior_drops = [d["where"] for d in cond["drop"] if d.get("where")]
            row_drops = drop_rows(working, prior_drops,
                                  {d["where"]: f"{d.get('label')} — {d['decision']}"
                                   for d in cond["drop"] if d.get("where")})
            row_drops["evaluated"] = cond
        to_pdf(working, pdf)
    except FillError as e:
        print(f"STOP: {e}", file=sys.stderr)
        return 2
    after_sha, _ = sha256_of(tmpl)
    if before_sha != after_sha:
        print("STOP: the TEMPLATE changed on disk during the fill. That must never happen.", file=sys.stderr)
        return 2

    # ---- the sendable class, and the scrub that follows from it.
    # Joe's doctrine (2026-07-31), which supersedes the old blanket "clients get
    # PDFs, never working files": a LETTER goes to the listing agent as a .docx
    # so they can edit and revise it, and that editing is the negotiation. A
    # SPREADSHEET goes as a PDF exactly so the formulas stay ours. So the
    # outbound artifact is per template kind, not one rule for both, and the
    # scrub applies to whichever one is actually going to leave.
    sendable_role = "working" if ext == ".docx" else "pdf"
    scrub = None
    if sendable_role == "working" and ext == ".docx":
        scrub = scrub_docx(working)
        to_pdf(working, pdf)          # re-render so the record copy matches the scrubbed file

    text = doc_text(working)
    leaks = leak_guard(text, plan.get("listing_side_names", []))
    lint = run_lint(text, plan["basename"][:40])
    color = color_check(working)

    # ---- the client copy (ORDER 23(c)), built from the FINISHED sendable so it
    # is the same letter Joe read with the audience rows removed. It runs the
    # same gates: an artifact that skips the gates is an artifact nobody checked.
    client = None
    if ext == ".docx" and not a.no_client_copy:
        try:
            client = build_client_copy(working, plan, fmap, prior_drops)
        except FillError as e:
            print(f"STOP: client copy: {e}", file=sys.stderr)
            return 2

    client_failed = bool(client and client.get("produced") and not client.get("passed"))
    blocked = bool(leaks) or lint["hard"] or not color["passed"] or client_failed
    routed_to, route_note = STAGING, "staging only (--route staging)"
    if a.route == "onedrive":
        if blocked:
            route_note = ("BLOCKED from OneDrive: "
                          + ("leak guard findings; " if leaks else "")
                          + ("writing-lint HARD findings; " if lint["hard"] else "")
                          + ("the colour/pipe finish gate failed; " if not color["passed"] else "")
                          + ("the CLIENT COPY failed its own gates" if client_failed else "")
                          + ". Staging copy only, per the ORDER 13 gate.")
        else:
            folder = find_deal_folder(plan)
            if not folder:
                route_note = ("no matching deal folder under Active Deals; staged instead of "
                              "creating a second folder beside the real one")
            else:
                filed = [working, pdf]
                if client and client.get("produced"):
                    filed += [client["docx"], client["pdf"]]
                for src in filed:
                    shutil.copy2(src, os.path.join(folder, os.path.basename(src)))
                routed_to, route_note = folder, "filed to the deal's existing OneDrive folder"

    out = {
        "document_id": plan["document_id"],
        "template": plan["template"]["slug"],
        "template_version": plan["template"]["template_version"],
        "template_unchanged": before_sha == after_sha,
        "working": working, "pdf": pdf,
        "working_sha256": sha256_of(working)[0], "working_bytes": sha256_of(working)[1],
        "pdf_sha256": sha256_of(pdf)[0], "pdf_bytes": sha256_of(pdf)[1],
        "slots_filled": len([e for e in plan["edits"] if not e.get("owed")]),
        "slots_owed": len(plan["owed"]),
        "slots_carried": len(plan["carried"]),
        "leak_findings": leaks,
        "lint_hard": lint["hard"], "lint_report": lint["report"],
        "sendable": {
            "role": sendable_role,
            "file": working if sendable_role == "working" else pdf,
            "not_sendable": pdf if sendable_role == "working" else working,
            "rule": ("Letters go to the listing agent as the WORKING .docx so they can edit and "
                     "revise it; the PDF is the record and preview copy."
                     if sendable_role == "working" else
                     "Spreadsheets go out as the PDF so the formulas stay ours; the working "
                     "workbook never leaves."),
            "human_gate": "Sendable names the format, not permission. Joe sends; no verb can."},
        "scrub": scrub,
        "color_finish": color_finish,
        "unresolved_options": options,
        "row_drops": row_drops,
        "row_drop_note": ("ORDER 23: the field map decides which rows leave a letter. `drop_when` "
                          "removes a term that does not exist on this deal (an NNN lease has no "
                          "base year); `audience` removes a term this reader has no business "
                          "seeing, on the CLIENT COPY only. Drops run after every fill, so no "
                          "address moves under a slot, and a row is never dropped on a guess."),
        "client_copy": client,
        "color_gate": color,
        "color_note": ("Joe's convention: blue or red marks text that gets replaced, and a letter "
                       "that goes out is all black. Every run in the sendable file is forced black "
                       "except the OWED markers, which stay red, so the only colour left is a "
                       "question a human still owes. What the template CARRIED is read off the "
                       "'carried' list and audit_template_colors.py, not off the document's ink. "
                       "This is a gate: coloured runs must equal OWED markers exactly, every marker "
                       "red, and no pipe anywhere in the body."),
        "routed_to": routed_to, "route_note": route_note,
        "status": "draft",
        "human_gate": "DRAFT for Joe to review. Nothing was sent; no verb in this system can send.",
    }

    if a.finish:
        out["finish"] = finish_records(plan, out, a.r2_bucket, a.r2_dry_run)
    print(json.dumps(out, indent=2))
    return 0


def finish_records(plan: dict, out: dict, bucket: str, dry_run: bool = False) -> dict:
    """Upload the archive copies, write the attachment rows, patch the document.

    Pipeline-writes-direct, the ORDER 17 precedent: this is file-side bookkeeping
    that no verb can perform, because the bytes only exist on this Mac. The
    document row itself was created by the verb under the full envelope.

    ORDER 20 closed the half ORDER 13 had to leave owed. `attachment.r2_key` is
    `not null unique`, so the upload is what makes the row insertable: the object
    goes up FIRST, under the quota guard, and only then does a row claim it. If
    the quota refuses, no row is written and the owed note says so in Joe's
    words rather than in an error code.
    """
    url = os.environ.get("DATABASE_URL")
    if not url:
        return {"skipped": "DATABASE_URL not set"}
    import psycopg
    res: dict[str, Any] = {"r2_bucket": bucket, "attachments": []}
    with psycopg.connect(url) as cn, cn.cursor() as cur:
        cur.execute("select id from actor where slug='joe'")
        actor_row = cur.fetchone()
        if actor_row is None:
            raise RuntimeError("actor target 'joe' was not found")
        actor = actor_row[0]

        cap, provenance = r2.quota_bytes(cn)
        led = r2.load_ledger(bucket)
        res["reconcile"] = r2.reconcile(led, cn, bucket)
        res["quota"] = {"cap_bytes": cap, "provenance": provenance}

        files = [("working", out["working"], out["working_sha256"], out["working_bytes"]),
                 ("pdf", out["pdf"], out["pdf_sha256"], out["pdf_bytes"])]
        att_ids: dict[str, Any] = {}
        refusal: Any = None
        for role, path, sha, size in files:
            key = r2.object_key(plan["deal"].get("client_ref"), sha, path)
            try:
                up = r2.upload(path, key, sha, size, cap, provenance, led, bucket, dry_run)
            except r2.QuotaExceeded as q:
                refusal = q
                res["quota_refusal"] = q.message
                res["quota_detail"] = q.detail
                break
            res["attachments"].append({"role": role, **up})
            if dry_run:
                continue
            # One object, one attachment row. A row that already claims this key
            # is reused rather than duplicated, which is what makes the rerun a
            # no-op instead of a unique-violation.
            cur.execute("select id from attachment where r2_key=%s", (key,))
            row = cur.fetchone()
            if row:
                att_ids[role] = row[0]
                continue
            cur.execute(
                """insert into attachment (subject_type, subject_id, r2_key, filename, mime,
                     sha256, bytes, created_by)
                   values (%s,%s,%s,%s,%s,%s,%s,%s) returning id""",
                ("deal", plan["deal"]["id"], key, os.path.basename(path),
                 r2.mime_for(path), sha, size, actor))
            attachment_row = cur.fetchone()
            if attachment_row is None:
                raise RuntimeError(f"attachment insert for {role} returned no id")
            att_ids[role] = attachment_row[0]

        if refusal is not None:
            note = (f"working: {out['working']} (sha256 {out['working_sha256'][:12]}) · "
                    f"pdf: {out['pdf']} (sha256 {out['pdf_sha256'][:12]}) · "
                    f"OWED: R2 archive copy refused by the self-enforced quota "
                    f"({refusal.detail['used_bytes']} of {refusal.detail['quota_bytes']} bytes used; "
                    f"this upload would have been over by {refusal.detail['over_by_bytes']}). "
                    f"The OneDrive copies stand; only the archive copy is missing.")
            cause = "files produced locally; R2 archive copy OWED (quota refusal)"
        else:
            note = (f"working: {out['working']} (sha256 {out['working_sha256'][:12]}) · "
                    f"pdf: {out['pdf']} (sha256 {out['pdf_sha256'][:12]}) · "
                    f"archived to R2 bucket {bucket}")
            cause = "files produced locally and archived to R2"

        if not dry_run:
            cur.execute(
                """update document set lint_passed=%s, leak_check_passed=%s, note=%s,
                     working_attachment=coalesce(%s, working_attachment),
                     pdf_attachment=coalesce(%s, pdf_attachment)
                   where id=%s""",
                (not out["lint_hard"], not out["leak_findings"], note,
                 att_ids.get("working"), att_ids.get("pdf"), plan["document_id"]))
            cur.execute(
                """insert into event (occurred_at, actor_id, verb, subject_type, subject_id, field,
                     new_value, cause, agent_rationale)
                   values (now(),%s,'prepare-document','deal',%s,'document.files',%s::jsonb,'automation_job',%s)""",
                (actor, plan["deal"]["id"],
                 json.dumps({"document_id": plan["document_id"],
                             "working_sha256": out["working_sha256"],
                             "pdf_sha256": out["pdf_sha256"],
                             "r2_keys": [a["key"] for a in res["attachments"]]}),
                 cause))
            cn.commit()
            res["document_patched"] = True
        res["attachments_owed"] = refusal is not None
        res["usage"] = r2.usage_summary(led, cap, provenance)
    return res


if __name__ == "__main__":
    raise SystemExit(main())
